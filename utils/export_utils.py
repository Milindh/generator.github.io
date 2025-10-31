"""
Export utilities for creating DOCX and PDF files
"""

from io import BytesIO

# Check for required libraries
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Install with: pip install python-docx")

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ fpdf2 not installed. Install with: pip install fpdf2")


def create_docx(cover_letter_text, contact_info=None):
    """
    Create a professionally formatted Word document with horizontal lines
    Matches the style: Name | Line | Contact | Line | Content
    
    Args:
        cover_letter_text (str): The cover letter content
        contact_info (dict, optional): Contact information to add as header
    
    Returns:
        BytesIO: Buffer containing the DOCX file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    print("📄 Creating DOCX file...")
    
    # Ensure name is after "Sincerely," if contact_info provided
    if contact_info and contact_info.get('name') and contact_info['name'] != 'Not Found':
        signature_name = contact_info['name']
        # Check if name is already after Sincerely
        if "Sincerely," in cover_letter_text:
            parts = cover_letter_text.split("Sincerely,")
            after_sincerely = parts[1].strip() if len(parts) > 1 else ""
            
            # If nothing after Sincerely or it's "[Your Name]", add the real name
            if not after_sincerely or after_sincerely == "[Your Name]":
                cover_letter_text = parts[0] + f"Sincerely,\n{signature_name}"
            elif signature_name not in after_sincerely:
                # Replace whatever is there with the extracted name
                cover_letter_text = parts[0] + f"Sincerely,\n{signature_name}"
    
    doc = Document()
    
    # Set margins (0.5 inch all around for tighter layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add contact header if provided
    if contact_info and contact_info.get('name') != 'Not Found':
        # Name - Centered, Bold, Large (18pt)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(contact_info['name'])
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(6)
        
        # First horizontal line (after name)
        line_para1 = doc.add_paragraph()
        line_para1.paragraph_format.space_before = Pt(0)
        line_para1.paragraph_format.space_after = Pt(6)
        line_run1 = line_para1.add_run('_' * 100)
        line_run1.font.name = 'Calibri'
        line_run1.font.size = Pt(11)
        line_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info - Centered, smaller (10pt)
        contact_parts = []
        if contact_info.get('phone') != 'Not Found':
            contact_parts.append(f"Phone: {contact_info['phone']}")
        if contact_info.get('email') != 'Not Found':
            contact_parts.append(contact_info['email'])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(' | '.join(contact_parts))
            contact_run.font.name = 'Calibri'
            contact_run.font.size = Pt(10)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(6)
        
        # Second horizontal line (after contact)
        line_para2 = doc.add_paragraph()
        line_para2.paragraph_format.space_before = Pt(0)
        line_para2.paragraph_format.space_after = Pt(12)
        line_run2 = line_para2.add_run('_' * 100)
        line_run2.font.name = 'Calibri'
        line_run2.font.size = Pt(11)
        line_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parse and add cover letter content - trust the Producer Agent's formatting
    lines = cover_letter_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines (we'll add them manually for spacing)
        if not line:
            doc.add_paragraph()  # Add blank line
            continue
        
        # Add the line as a paragraph
        para = doc.add_paragraph(line)
        
        # Style based on content
        if line.startswith('Dear') or line in ['Sincerely,', 'Best regards,', 'Regards,', 'Thank you,']:
            # Salutation and closing - left aligned
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Body paragraphs - justified
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 1.15
        
        # Font formatting
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    print("✓ DOCX created")
    
    return buffer


def create_pdf(cover_letter_text, contact_info=None):
    """
    Create a professionally formatted PDF document with horizontal lines
    Matches the style: Name | Line | Contact | Line | Content
    
    Args:
        cover_letter_text (str): The cover letter content
        contact_info (dict, optional): Contact information to add as header
    
    Returns:
        bytes: PDF file as bytes
    """
    if not PDF_AVAILABLE:
        raise ImportError("fpdf2 not installed. Run: pip install fpdf2")
    
    print("📕 Creating PDF file...")
    
    # Clean text for PDF - replace special characters that cause issues
    cover_letter_text = cover_letter_text.replace('\u2013', '-')  # en-dash to hyphen
    cover_letter_text = cover_letter_text.replace('\u2014', '-')  # em-dash to hyphen
    cover_letter_text = cover_letter_text.replace('\u2018', "'")  # left single quote
    cover_letter_text = cover_letter_text.replace('\u2019', "'")  # right single quote
    cover_letter_text = cover_letter_text.replace('\u201c', '"')  # left double quote
    cover_letter_text = cover_letter_text.replace('\u201d', '"')  # right double quote
    cover_letter_text = cover_letter_text.replace('\u2026', '...')  # ellipsis
    
    # Ensure name is after "Sincerely," if contact_info provided
    if contact_info and contact_info.get('name') and contact_info['name'] != 'Not Found':
        signature_name = contact_info['name']
        # Check if name is already after Sincerely
        if "Sincerely," in cover_letter_text:
            parts = cover_letter_text.split("Sincerely,")
            after_sincerely = parts[1].strip() if len(parts) > 1 else ""
            
            # If nothing after Sincerely or it's "[Your Name]", add the real name
            if not after_sincerely or after_sincerely == "[Your Name]":
                cover_letter_text = parts[0] + f"Sincerely,\n{signature_name}"
            elif signature_name not in after_sincerely:
                # Replace whatever is there with the extracted name
                cover_letter_text = parts[0] + f"Sincerely,\n{signature_name}"
    
    pdf = FPDF()
    pdf.add_page()
    
    # Set margins
    pdf.set_margins(left=20, top=15, right=20)
    
    # Add contact header if provided
    if contact_info and contact_info.get('name') != 'Not Found':
        # Name - Centered, Bold, Large (18pt)
        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 10, contact_info['name'], ln=True, align='C')
        pdf.ln(2)
        
        # First horizontal line (after name)
        pdf.set_draw_color(0, 0, 0)  # Black color
        pdf.set_line_width(0.5)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())  # Draw line across page
        pdf.ln(5)
        
        # Contact info - Centered, smaller (10pt)
        pdf.set_font('Arial', '', 10)
        contact_parts = []
        if contact_info.get('phone') != 'Not Found':
            contact_parts.append(f"Phone: {contact_info['phone']}")
        if contact_info.get('email') != 'Not Found':
            contact_parts.append(contact_info['email'])
        
        if contact_parts:
            pdf.cell(0, 6, ' | '.join(contact_parts), ln=True, align='C')
        
        # Second horizontal line (after contact)
        pdf.ln(2)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())  # Draw line across page
        pdf.ln(8)
    
    # Process cover letter content - trust the Producer Agent's formatting
    pdf.set_font('Arial', '', 11)
    
    lines = cover_letter_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines but add spacing
        if not line:
            pdf.ln(6)
            continue
        
        # Check if special line (salutation/closing)
        if line.startswith('Dear') or line in ['Sincerely,', 'Best regards,', 'Regards,', 'Thank you,']:
            pdf.cell(0, 6, line, ln=True, align='L')
        else:
            # Body paragraph - justified
            pdf.multi_cell(0, 6, line, align='J')
    
    # Return PDF as bytes
    pdf_output = pdf.output(dest='S')
    
    # Convert to bytes if it's a string
    if isinstance(pdf_output, str):
        pdf_bytes = pdf_output.encode('latin-1')
    else:
        pdf_bytes = pdf_output
    
    print("✓ PDF created")
    
    return pdf_bytes


# Test the export functions
if __name__ == "__main__":
    test_letter = """Dear Hiring Manager,

I am excited to apply for this position. I have 5 years of experience in software development.

In my current role, I lead a team of developers and work on cloud infrastructure.

Sincerely,
John Smith"""
    
    test_contact = {
        'name': 'John Smith',
        'email': 'john.smith@email.com',
        'phone': '(555) 123-4567',
        'address': 'San Francisco, CA'
    }
    
    print("Testing export utilities...")
    print("="*50)
    
    # Test DOCX
    if DOCX_AVAILABLE:
        try:
            docx_buffer = create_docx(test_letter, test_contact)
            print(f"✓ DOCX test passed - {len(docx_buffer.getvalue())} bytes")
        except Exception as e:
            print(f"✗ DOCX test failed: {e}")
    
    # Test PDF
    if PDF_AVAILABLE:
        try:
            pdf_bytes = create_pdf(test_letter, test_contact)
            print(f"✓ PDF test passed - {len(pdf_bytes)} bytes")
        except Exception as e:
            print(f"✗ PDF test failed: {e}")